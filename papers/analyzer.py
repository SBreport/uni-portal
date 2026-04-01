"""
논문 분석 프로그램 — uni-portal papers DB 연동

미용의료 논문 PDF를 Claude API로 분석하여:
1. 의학 비전문가도 이해할 수 있는 서술형 요약 생성
2. uni-portal papers DB에 저장 (기본)
3. Word(.docx) / Excel(.xlsx) 문서로 출력 (선택)

사용법:
  python papers/analyzer.py sample.pdf                 # 분석 → DB 저장
  python papers/analyzer.py --dir pdfs/                # 폴더 내 전체 분석
  python papers/analyzer.py --dry-run sample.pdf       # 분석만 (DB 저장 안 함)
  python papers/analyzer.py --export-docx sample.pdf   # 분석 + DB + Word 출력
  python papers/analyzer.py --export-xlsx sample.pdf   # 분석 + DB + Excel 출력
  python papers/analyzer.py --export-only results.json # 기존 JSON → 문서 출력
"""

import os
import re
import sys
import json
import time
import hashlib
import argparse
import sqlite3

# 프로젝트 루트를 sys.path에 추가 (equipment.matcher 등 import용)
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT_FOR_PATH = os.path.dirname(_SCRIPT_DIR)
if _PROJECT_ROOT_FOR_PATH not in sys.path:
    sys.path.insert(0, _PROJECT_ROOT_FOR_PATH)
from pathlib import Path
from datetime import datetime
from shared.db import now_str

import fitz  # PyMuPDF
import requests
import anthropic
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

# ── 설정 ──
API_BASE = "http://localhost:8000"
_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DB_PATH = os.path.join(_PROJECT_ROOT, "data", "equipment.db")
RESULTS_DIR = os.path.join(_PROJECT_ROOT, "paper_results")
CLAUDE_MODEL = "claude-sonnet-4-20250514"

EVIDENCE_STARS = {0: "☆☆☆☆☆", 1: "★☆☆☆☆", 2: "★★☆☆☆", 3: "★★★☆☆", 4: "★★★★☆", 5: "★★★★★"}

ANALYSIS_PROMPT = """당신은 미용의료 논문을 분석하여, 의학 지식이 전혀 없는 칼럼 작가가 이해할 수 있도록 쉽고 상세하게 요약하는 전문가입니다.

## 분석 대상 논문 텍스트:
{paper_text}

## 출력 형식 (JSON):
반드시 아래 JSON 형식으로만 응답하세요. JSON 외의 텍스트는 포함하지 마세요.

{{
  "title": "논문의 영문 원제목",
  "title_ko": "한국어 제목 (아래 번역 원칙 참고)",
  "authors": "저자 (첫 저자 et al. 형태)",
  "journal": "학술지명",
  "pub_year": 출판연도(숫자) 또는 null,
  "one_line_summary": "육하원칙 기반 한 줄 요약 (아래 작성 원칙 참고)",
  "key_results_brief": "핵심 결과를 글머리 기호로 정리. 각 항목에 구체적 수치 포함. 예: - 피부 탄력 32% 개선\\n- 환자 만족도 87%",
  "research_purpose": "연구 목적을 서술형 줄 글로 상세하게 설명 (3~5문장). 이 연구를 왜 했는지, 무엇을 확인하려 했는지를 의학 비전문가도 이해할 수 있게 풀어서 작성. 전문 용어가 나오면 반드시 괄호 안에 쉬운 설명을 추가. 복잡한 개념은 '※ 쉽게 말하면:' 으로 한 번 더 풀어서 설명.",
  "study_design": "연구 설계 및 방법을 서술형 줄 글로 상세하게 설명 (5~10문장). 연구 유형, 대상자 수와 특성, 시술 방법, 측정 방법, 추적 기간 등을 구체적으로 서술. 전문 용어는 괄호 안에 쉬운 말로 풀이. 예: '안면거상술(rhytidectomy, 얼굴 주름 제거 수술) 예정 여성 2명이었습니다.'",
  "key_results": "핵심 결과를 서술형 줄 글로 상세하게 설명 (5~10문장). 구체적인 수치(백분율, p-value, 효과 크기 등)를 반드시 포함. 결과의 의미를 쉬운 말로 해석. 예: '시술 부위의 콜라겐 생성률이 비시술 부위보다 평균 42% 높았다고 보고했습니다.'",
  "conclusion": "연구 결론을 서술형 줄 글로 설명 (3~5문장). 저자들이 어떤 결론을 내렸는지, 이것이 실제 시술에 어떤 의미가 있는지를 쉬운 말로 정리.",
  "cautions": "해석 시 주의점. 연구의 한계를 솔직하게 기재 (글머리 기호 사용). 반드시 포함할 것: 표본 크기의 적정성, 연구 설계의 한계(무작위배정 여부 등), 추적 기간의 충분성, 이해상충(연구 스폰서/저자 소속이 해당 시술 제조사인 경우 명시), 통계적 유의성 여부.",
  "keywords": ["한국어 키워드", "시술명 포함"],
  "evidence_level": 0에서 5 사이 정수,
  "evidence_level_reason": "근거 수준을 이 값으로 판정한 이유를 1~2문장으로 설명",
  "study_type": "메타분석|체계적문헌고찰|RCT|비무작위대조시험|전향적코호트|후향적코호트|환자대조연구|증례시리즈|증례보고|전문가리뷰|내러티브리뷰|전문가합의(Delphi)|기초연구(in vitro/ex vivo)|동물실험 중 택1",
  "sample_size": "대상자 수 (예: '120명')",
  "follow_up_period": "추적 관찰 기간 (예: '6개월')",
  "related_devices": ["논문에서 다루는 장비/시술 브랜드명 — 아래 '등록 장비 목록'에 있는 이름을 우선 사용. 매칭되는 장비가 없으면 논문에 나온 원래 장비명을 기재. 특정 기술(HIFU, RF, IPL, KTP 등)의 원리/기전을 다루는 논문이면 기술명도 함께 포함 (예: ['울쎄라', 'HIFU'])"],
  "is_aesthetic_medical": true 또는 false,
  "non_aesthetic_reason": "미용의료와 무관한 경우 그 이유를 한 줄로 설명. 관련 있으면 빈 문자열"
}}

## 핵심 작성 원칙:

### 제목 번역 원칙 (title_ko):
- **원문 의미를 최대한 보존**하면서, 자연스러운 한국어로 번역. 과도한 함축이나 의역 금지.
- **고정 구조**: [시술명/관련 시술] + [기술 설명] + [적응증/효과] + [문헌 형식]
- 시술명이 명확한 경우(울쎄라, 써마지 등) 제목 앞에 시술명 배치.
- 기술명(MFU-V, RF 등)은 그대로 유지하되, 한국어 풀이를 앞이나 뒤에 추가.
- 어색한 학술 번역어만 교체: "회춘" → "노화 개선" 또는 "피부 개선", "비침습적" → "비수술적"
- 예시:
  - 원문: "Micro-focused Ultrasound with Visualization (MFU-V) for Non-invasive Facial Rejuvenation: A Comprehensive Review"
  - BAD (과도한 함축): "울쎄라 리프팅의 효과와 안전성 — 임상 연구 종합 분석"
  - BAD (어색한 직역): "비침습적 얼굴 회춘을 위한 시각화 마이크로 집속 초음파(MFU-V): 종합적 리뷰"
  - GOOD: "울쎄라에 적용되는 시각화 마이크로 집속 초음파(MFU-V)의 안면 노화 개선 효과: 종합 리뷰"
  - GOOD: "울쎄라 관련 초음파 리프팅 기술(MFU-V)의 얼굴 리프팅·탄력 개선 효과: 종합 리뷰"

### 본문 번역 원칙 (research_purpose, study_design, key_results, conclusion 등):
- 원문의 의미와 맥락을 충실히 전달하되, 비전문가가 읽기 어려운 표현만 자연스럽게 교체.
- 전문 용어는 제거하지 말고, 처음 등장 시 괄호 안에 쉬운 설명을 추가하는 방식 유지.
- 과도한 의역이나 원문에 없는 해석 추가 금지.

### 한 줄 요약 원칙 (one_line_summary):
- **육하원칙**에 따라 구조화하여 작성:
  - 누가(연구진/기관): 어디서 수행한 연구인지
  - 무엇을(대상): 어떤 시술/장비를 연구했는지
  - 어떻게(방법): 몇 명을 대상으로, 어떤 방식으로 연구했는지
  - 왜(목적): 무엇을 확인하려 했는지
  - 결과: 핵심 수치와 함께 어떤 결과가 나왔는지
- 3~4문장으로 작성. 이 요약만 읽으면 논문의 핵심을 파악할 수 있어야 함.
- 예시:
  - BAD: "이 연구는 울쎄라 시술의 효과를 확인했습니다."
  - GOOD: "울쎄라(초음파 리프팅) 시술의 효과와 안전성을 검증하기 위해, 경도~중등도 얼굴 처짐이 있는 환자들을 대상으로 다수의 임상 연구를 종합 분석했습니다. 시술 3개월 후 60~70%의 환자에서 리프팅 효과가 확인되었고, 효과는 평균 6~12개월 지속되었습니다. 부작용은 대부분 일시적인 통증과 붓기로, 수일 내 자연 해소되었습니다."

### 서술 원칙:
1. **서술형 줄 글**: research_purpose, study_design, key_results, conclusion은 반드시 완성된 문장으로 충분히 풀어서 작성.
2. **전문 용어 풀이**: 의학/과학 전문 용어가 처음 나올 때 괄호 안에 쉬운 한국어 설명 추가. 예: "SMAS층(피부 아래 근막층, 리프팅 시술의 핵심 타겟)"
3. **보충 설명**: 특히 복잡한 실험 방법이나 개념은 "※ 쉽게 말하면:" 형태로 한 번 더 풀어서 설명
4. **구체적 수치**: 가능한 모든 수치 데이터(%, p-value, 인원수, 기간) 포함
5. **이해상충 투명 공개**: 연구 스폰서가 시술 장비 제조사인 경우 반드시 cautions에 명시
6. **근거 수준 기준**: 5=메타분석/체계적문헌고찰, 4=RCT, 3=코호트/비교연구, 2=증례보고/시리즈, 1=전문가의견/리뷰/전문가합의, 0=미분류 또는 기초연구
7. **study_type 선택지**: 메타분석, 체계적문헌고찰, RCT, 비무작위대조시험, 전향적코호트, 후향적코호트, 환자대조연구, 증례시리즈, 증례보고, 전문가리뷰, 내러티브리뷰, **전문가합의(Delphi)**, **기초연구(in vitro/ex vivo)**, **동물실험** 중 택1

### 미용의료 관련성 판단 (is_aesthetic_medical):
- 미용의료 범위: 피부 시술(레이저, 초음파, 고주파, 필러, 보톡스, 스킨부스터, 탈모, 제모, 여드름, 색소, 흉터, 비만치료 등)
- 미용의료와 **무관**한 논문 예: 순수 치과학, 내과, 외과, 소아과, 안과(미용 제외) 등
- 논문이 미용시술 장비의 **원리·기전** 연구이면 관련 있음으로 판단 (예: HIFU 조직 응고 기전 = 관련)
- 판단이 애매한 경우 true로 설정

### 등록 장비 목록 (related_devices 매칭에 사용):
아래 목록에서 논문 내용과 가장 관련 있는 장비명을 찾아 related_devices에 기재하세요.
장비명은 **목록에 있는 이름 그대로** 사용하세요. (예: "Thermage FLX" → "써마지FLX")
한 논문에 여러 장비가 관련되면 모두 포함하세요.

{device_list_hint}
"""


# ── PDF 텍스트 추출 ──

def extract_text_from_pdf(pdf_path: str) -> str:
    """PDF에서 텍스트 추출 (PyMuPDF)."""
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text


def extract_doi(text: str) -> str:
    """텍스트에서 DOI 추출."""
    match = re.search(r'(10\.\d{4,}/[^\s\]>]+)', text)
    if match:
        doi = match.group(1).rstrip('.,;:)')
        return doi
    return ""


# ── CrossRef 메타데이터 ──

def fetch_crossref_metadata(doi: str) -> dict:
    """CrossRef API로 논문 메타데이터 조회."""
    if not doi:
        return {}
    try:
        resp = requests.get(
            f"https://api.crossref.org/works/{doi}",
            headers={"User-Agent": "UniPortal/1.0 (paper-analyzer)"},
            timeout=10,
        )
        if resp.ok:
            data = resp.json()["message"]
            authors_list = data.get("author", [])
            if authors_list:
                first = authors_list[0]
                authors = f"{first.get('family', '')} {first.get('given', '')}"
                if len(authors_list) > 1:
                    authors += " et al."
            else:
                authors = ""
            pub_parts = data.get("published-print", data.get("published-online", {}))
            date_parts = pub_parts.get("date-parts", [[None]])[0] if pub_parts else [None]
            return {
                "title": data.get("title", [""])[0],
                "authors": authors,
                "journal": data.get("container-title", [""])[0],
                "pub_year": date_parts[0] if date_parts else None,
            }
    except Exception as e:
        print(f"  [CrossRef 조회 실패: {e}]")
    return {}


# ── Claude API 분석 ──

def build_device_hint(device_list: list[dict]) -> str:
    """DB 장비 목록에서 Claude에게 제공할 힌트 문자열 생성."""
    lines = []
    for d in device_list:
        aliases = d.get("aliases", "")
        category = d.get("category", "")
        name = d["name"]
        parts = [name]
        if aliases:
            parts.append(f"(별칭: {aliases})")
        if category:
            parts.append(f"[{category}]")
        lines.append(" ".join(parts))
    return "\n".join(lines)


def analyze_with_claude(text: str, device_list: list[dict] = None) -> dict:
    """Claude API로 논문 분석."""
    client = anthropic.Anthropic()

    # 토큰 제한: 텍스트가 너무 길면 앞뒤 잘라서 전달
    max_chars = 80000
    if len(text) > max_chars:
        half = max_chars // 2
        text = text[:half] + "\n\n[... 중간 생략 ...]\n\n" + text[-half:]

    # 장비 목록 힌트 구성
    device_hint = ""
    if device_list:
        device_hint = build_device_hint(device_list)

    prompt = ANALYSIS_PROMPT.format(paper_text=text, device_list_hint=device_hint)

    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )

    response_text = message.content[0].text.strip()

    # JSON 파싱 (코드 블록 제거)
    if response_text.startswith("```"):
        response_text = re.sub(r'^```(?:json)?\s*', '', response_text)
        response_text = re.sub(r'\s*```$', '', response_text)

    try:
        return json.loads(response_text)
    except json.JSONDecodeError as e:
        print(f"  [JSON 파싱 실패: {e}]")
        print(f"  [응답 일부: {response_text[:200]}...]")
        return {}


# ── 장비/시술 매칭 ──

def get_device_info_list() -> list[dict]:
    """DB에서 device_info 목록 조회."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    rows = conn.execute("SELECT * FROM device_info ORDER BY usage_count DESC, name").fetchall()
    conn.close()
    return [dict(r) for r in rows]


def get_treatments_list() -> list[dict]:
    """DB에서 시술 목록 조회."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    rows = conn.execute("SELECT id, name, brand, category_id FROM evt_treatments").fetchall()
    conn.close()
    return [dict(r) for r in rows]


def match_device_id(related_devices: list[str], device_list: list[dict]) -> tuple:
    """장비명으로 device_info_id 매칭 (단일, 하위호환). 공용 matcher 엔진 사용."""
    from equipment.matcher import match_from_names
    return match_from_names(related_devices, device_list)


def match_devices_multi(related_devices: list[str], device_list: list[dict]) -> list[dict]:
    """장비명으로 다대다 매칭 (직접 + 기술). paper_devices 테이블용."""
    from equipment.matcher import match_devices_for_paper
    return match_devices_for_paper(related_devices, device_list)


def match_treatment_id(related_devices: list[str], treatment_list: list[dict]) -> tuple:
    """시술명으로 treatment_id 매칭. 공용 matcher 엔진 사용."""
    from equipment.matcher import match_treatment_from_names
    return match_treatment_from_names(related_devices, treatment_list)


# ── DB 마이그레이션 ──

def ensure_new_columns():
    """기존 papers 테이블에 새 컬럼이 없으면 추가."""
    new_cols = [
        ("one_line_summary", "TEXT DEFAULT ''"),
        ("research_purpose", "TEXT DEFAULT ''"),
        ("study_design_detail", "TEXT DEFAULT ''"),
        ("key_results", "TEXT DEFAULT ''"),
        ("conclusion", "TEXT DEFAULT ''"),
        ("quotable_stats", "TEXT DEFAULT '[]'"),
        ("cautions", "TEXT DEFAULT ''"),
        ("follow_up_period", "TEXT DEFAULT ''"),
        ("file_hash", "TEXT DEFAULT ''"),
    ]
    conn = sqlite3.connect(DB_PATH)
    existing = {row[1] for row in conn.execute("PRAGMA table_info(papers)").fetchall()}
    for col_name, col_def in new_cols:
        if col_name not in existing:
            conn.execute(f"ALTER TABLE papers ADD COLUMN {col_name} {col_def}")
            print(f"  [DB 마이그레이션] papers.{col_name} 컬럼 추가")
    conn.commit()
    conn.close()


# ── 학술 논문 검증 ──

# 학술 논문에서 흔히 발견되는 섹션 키워드
ACADEMIC_MARKERS = [
    # 영문 학술 구조
    r'\babstract\b', r'\bintroduction\b', r'\bmethods?\b', r'\bresults?\b',
    r'\bdiscussion\b', r'\bconclusions?\b', r'\breferences\b', r'\bbibliography\b',
    r'\backnowledg', r'\bfunding\b', r'\bconflict.{0,5}interest',
    # 저널/학술 표기
    r'\boriginal\s+article\b', r'\bresearch\s+article\b', r'\breview\s+article\b',
    r'\bcase\s+report\b', r'\bclinical\s+trial\b',
    r'\bjournal\s+of\b', r'\bdoi[:\s]', r'\bISSN\b', r'\bPMID\b', r'\bPubMed\b',
    # 학술 통계 표현
    r'p\s*[<>=]\s*0\.\d', r'\bCI\s*[:=]', r'\bn\s*=\s*\d',
    # 한국어 학술 구조
    r'초록', r'서론', r'연구\s*방법', r'연구\s*결과', r'고찰', r'참고\s*문헌',
]

# 비학술 자료 (AI 생성, 개인 정리 등)의 특징
NON_ACADEMIC_MARKERS = [
    r'\[[\d]{1,3}\].*\[[\d]{1,3}\]',  # [1] [2] [3] 형태 참조 번호 반복
    r'쉽게\s*말하면', r'정리하면',  # 구어체 설명
    r'~라고\s*(생각|보면)', r'~(이|가)\s*핵심이다',  # 구어체
    r'perplexity|퍼플렉시티|chatgpt|GPT|Claude|AI\s*(가|에|로)',  # AI 도구 언급
    r'조사해\s*줘', r'알려\s*줘', r'정리해\s*줘',  # AI 프롬프트 잔존
]


def validate_academic_paper(text: str, pdf_path: str) -> dict:
    """PDF가 학술 논문인지 검증.

    Returns:
        {"valid": True} 또는
        {"valid": False, "reason": "...", "score": n, "details": [...]}
    """
    text_lower = text.lower()
    head_text = text[:3000]  # 앞부분에 메타정보가 집중
    details = []

    # 1) 텍스트 길이 체크 — 학술 논문은 보통 5000자 이상
    if len(text) < 3000:
        details.append(f"텍스트가 매우 짧음 ({len(text):,}자, 학술 논문은 보통 5,000자 이상)")

    # 2) 학술 마커 점수 (높을수록 학술적)
    academic_score = 0
    for pattern in ACADEMIC_MARKERS:
        if re.search(pattern, text_lower):
            academic_score += 1

    # 3) 비학술 마커 점수 (높을수록 비학술적)
    non_academic_score = 0
    non_academic_found = []
    for pattern in NON_ACADEMIC_MARKERS:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            non_academic_score += len(matches)
            non_academic_found.append(f"'{pattern}' 패턴 {len(matches)}회")

    # 4) 저자·소속 정보 존재 여부
    has_author_info = bool(re.search(
        r'(MD|PhD|MS|BSc|Department|University|Hospital|Clinic|Institute|의과대학|병원|연구소)',
        head_text, re.IGNORECASE
    ))

    # 5) DOI 존재 여부
    has_doi = bool(re.search(r'10\.\d{4,}/[^\s]+', text))

    # 6) 참고문헌 섹션 존재 여부
    has_references = bool(re.search(
        r'(references|bibliography|참고\s*문헌|인용)',
        text_lower
    ))

    # 판정 로직
    # 학술 마커 3개 미만 + 비학술 마커 2개 이상 → 비학술
    # 학술 마커 3개 미만 + 저자 정보 없음 + DOI 없음 → 비학술
    is_valid = True
    reason_parts = []

    if non_academic_score >= 3:
        details.append(f"비학술 패턴 다수 발견 ({non_academic_score}건)")
        is_valid = False

    if academic_score < 3 and not has_doi and not has_author_info:
        details.append(f"학술 구조 부족 (학술 마커 {academic_score}개, DOI 없음, 저자 정보 없음)")
        is_valid = False

    if academic_score < 2 and non_academic_score >= 2:
        details.append("학술 마커 거의 없고 비학술 패턴 다수")
        is_valid = False

    if len(text) < 2000 and academic_score < 3:
        details.append("텍스트 매우 짧고 학술 구조 부족")
        is_valid = False

    if is_valid:
        return {"valid": True, "academic_score": academic_score}

    return {
        "valid": False,
        "reason": " / ".join(details),
        "academic_score": academic_score,
        "non_academic_score": non_academic_score,
        "has_doi": has_doi,
        "has_author_info": has_author_info,
        "has_references": has_references,
        "details": details,
    }


# ── 중복 체크 ──

def compute_file_hash(pdf_path: str) -> str:
    """PDF 파일의 SHA-256 해시 계산."""
    h = hashlib.sha256()
    with open(pdf_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def normalize_title(title: str) -> str:
    """제목 비교용 정규화: 소문자, 특수문자/공백 제거."""
    title = title.lower().strip()
    title = re.sub(r'[^a-z0-9가-힣]', '', title)
    return title


def check_duplicate(pdf_path: str, text: str, doi: str) -> dict | None:
    """PDF가 DB에 이미 등록된 논문인지 확인.

    체크 순서:
    1. 파일 해시 — 동일 파일 재분석 방지
    2. DOI 일치 — 가장 확실한 중복 판별
    3. 제목 유사도 — PDF 텍스트 앞부분과 DB 제목 비교

    Returns: 중복 논문 dict (id, title, 사유) 또는 None
    """
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row

    # 1) 파일 해시 체크
    file_hash = compute_file_hash(pdf_path)
    row = conn.execute(
        "SELECT id, title, title_ko FROM papers WHERE file_hash = ? AND status != 'deleted'",
        (file_hash,)
    ).fetchone()
    if row:
        conn.close()
        return {
            "id": row["id"],
            "title": row["title_ko"] or row["title"],
            "reason": "동일 파일 (파일 해시 일치)",
            "file_hash": file_hash,
        }

    # 2) DOI 체크
    if doi:
        # DOI 끝의 .pdf 등 불필요한 확장자 제거
        clean_doi = re.sub(r'\.(pdf|html|xml)$', '', doi, flags=re.IGNORECASE)
        row = conn.execute(
            "SELECT id, title, title_ko FROM papers WHERE (doi = ? OR doi = ?) AND status != 'deleted'",
            (doi, clean_doi)
        ).fetchone()
        if row:
            conn.close()
            return {
                "id": row["id"],
                "title": row["title_ko"] or row["title"],
                "reason": f"동일 DOI ({clean_doi})",
                "file_hash": file_hash,
            }

    # 3) 제목 유사도 체크 — PDF 텍스트 앞 2000자에서 긴 문장을 추출해 DB 제목과 비교
    all_papers = conn.execute(
        "SELECT id, title, title_ko FROM papers WHERE status != 'deleted'"
    ).fetchall()
    conn.close()

    if all_papers:
        # PDF 앞부분에서 제목 후보 추출 (첫 2000자 내 가장 긴 줄들)
        head_text = text[:2000]
        head_lines = [ln.strip() for ln in head_text.split('\n') if len(ln.strip()) > 10]

        for paper in all_papers:
            db_title_norm = normalize_title(paper["title"])
            db_title_ko_norm = normalize_title(paper["title_ko"] or "")

            if not db_title_norm:
                continue

            # PDF 텍스트 앞부분에 DB 제목이 포함되어 있는지
            head_norm = normalize_title(head_text)
            if len(db_title_norm) >= 15 and db_title_norm in head_norm:
                return {
                    "id": paper["id"],
                    "title": paper["title_ko"] or paper["title"],
                    "reason": "제목 일치 (PDF 텍스트에서 기존 논문 제목 발견)",
                    "file_hash": file_hash,
                }

    return None  # 중복 아님


# ── PDF 1건 처리 ──

def process_pdf(pdf_path: str, device_list: list[dict], treatment_list: list[dict]) -> dict:
    """PDF 1건을 분석하여 papers 데이터로 변환."""
    print(f"\n{'='*60}")
    print(f"분석 중: {pdf_path}")
    print(f"{'='*60}")

    # 1) 텍스트 추출
    text = extract_text_from_pdf(pdf_path)
    if not text.strip():
        print("  [오류] PDF에서 텍스트를 추출할 수 없습니다.")
        return None

    print(f"  텍스트 추출: {len(text):,}자")

    # 2) DOI 추출
    doi = extract_doi(text)
    if doi:
        print(f"  DOI: {doi}")

    # 3) 학술 논문 검증 (API 호출 전에 수행)
    validation = validate_academic_paper(text, pdf_path)
    if not validation["valid"]:
        print(f"  [비학술 자료] 학술 논문이 아닌 것으로 판단되어 건너뜁니다.")
        print(f"    사유: {validation['reason']}")
        print(f"    (학술 마커: {validation['academic_score']}개, 비학술 패턴: {validation['non_academic_score']}건)")
        return {"_skipped": True, "_reason": f"비학술 자료 — {validation['reason']}", "_file": str(pdf_path)}
    else:
        print(f"  학술 검증: 통과 (학술 마커 {validation['academic_score']}개)")

    # 4) 중복 체크 (API 호출 전에 수행)
    dup = check_duplicate(pdf_path, text, doi)
    if dup:
        print(f"  [중복 감지] 이미 등록된 논문입니다!")
        print(f"    사유: {dup['reason']}")
        print(f"    기존 논문: #{dup['id']} - {dup['title']}")
        return {"_skipped": True, "_reason": dup["reason"], "_existing_id": dup["id"], "_existing_title": dup["title"], "_file": str(pdf_path)}

    # 5) CrossRef 메타데이터
    meta = fetch_crossref_metadata(doi)
    if meta:
        print(f"  CrossRef: {meta.get('title', '')[:50]}...")

    # 5) Claude API 분석
    print("  Claude API 분석 중...")
    analysis = analyze_with_claude(text, device_list)
    if not analysis:
        print("  [오류] Claude API 분석 실패")
        return None

    print(f"  분석 완료: {analysis.get('title_ko', '')[:40]}...")

    # 5-1) 미용의료 관련성 체크
    is_aesthetic = analysis.get("is_aesthetic_medical", True)
    non_aesthetic_reason = analysis.get("non_aesthetic_reason", "")
    if not is_aesthetic:
        print(f"  ⚠ [비관련 논문] 미용의료와 무관: {non_aesthetic_reason}")
        # 비관련이어도 분석 결과는 반환 (사용자가 판단)
        analysis["_non_aesthetic"] = True
        analysis["_non_aesthetic_reason"] = non_aesthetic_reason

    # 6) 장비/시술 매칭 (다대다)
    related = analysis.get("related_devices", [])
    device_matches = match_devices_multi(related, device_list)
    treatment_id, treatment_name = match_treatment_id(related, treatment_list)

    # 하위호환: 첫 번째 직접 매칭을 대표 장비로 설정
    direct_matches = [m for m in device_matches if m["match_type"] == "direct"]
    tech_matches = [m for m in device_matches if m["match_type"] == "technology"]
    device_id = direct_matches[0]["device_info_id"] if direct_matches else None
    device_name = direct_matches[0]["device_name"] if direct_matches else ""

    if device_matches:
        direct_names = [m["device_name"] for m in direct_matches]
        tech_names = [m["device_name"] for m in tech_matches]
        if direct_names:
            print(f"  장비 매칭(직접): {', '.join(direct_names)}")
        if tech_names:
            print(f"  장비 매칭(기술): {', '.join(tech_names)} ← {set(m['match_keyword'] for m in tech_matches)}")
    else:
        print(f"  장비 매칭: 없음 (관련 장비: {related})")

    # 6) 데이터 조합
    result = {
        "device_info_id": device_id,
        "treatment_id": treatment_id,
        "device_name": device_name or "",
        "treatment_name": treatment_name or "",
        "doi": doi,
        "title": meta.get("title") or analysis.get("title", ""),
        "title_ko": analysis.get("title_ko", ""),
        "authors": meta.get("authors") or analysis.get("authors", ""),
        "journal": meta.get("journal") or analysis.get("journal", ""),
        "pub_year": meta.get("pub_year") or analysis.get("pub_year"),
        "pub_date": "",
        "abstract_summary": analysis.get("one_line_summary", ""),
        "key_findings": analysis.get("key_results_brief", ""),
        "keywords": json.dumps(analysis.get("keywords", []), ensure_ascii=False),
        "evidence_level": analysis.get("evidence_level", 0),
        "study_type": analysis.get("study_type", ""),
        "sample_size": analysis.get("sample_size", ""),
        "source_url": f"https://doi.org/{doi}" if doi else "",
        "source_file": str(pdf_path),
        "status": "draft",
        "one_line_summary": analysis.get("one_line_summary", ""),
        "research_purpose": analysis.get("research_purpose", ""),
        "study_design_detail": analysis.get("study_design", ""),
        "key_results": analysis.get("key_results", ""),
        "conclusion": analysis.get("conclusion", ""),
        "quotable_stats": json.dumps(
            analysis.get("key_results_brief", "").split("\n") if isinstance(analysis.get("key_results_brief"), str) else [],
            ensure_ascii=False,
        ),
        "cautions": "\n".join(analysis["cautions"]) if isinstance(analysis.get("cautions"), list) else analysis.get("cautions", ""),
        "follow_up_period": analysis.get("follow_up_period", ""),
        "file_hash": compute_file_hash(pdf_path),
        # 분석 메타 (문서 출력용, DB 저장 안 함)
        "_evidence_reason": analysis.get("evidence_level_reason", ""),
        "_related_devices": related,
        "_related_devices_raw": related,
        "_device_matches": device_matches,  # 다대다 매칭 결과
        "_non_aesthetic": analysis.get("_non_aesthetic", False),
        "_non_aesthetic_reason": analysis.get("_non_aesthetic_reason", ""),
    }

    return result


# ── DB 저장 ──

def save_to_db(papers: list[dict]):
    """papers DB에 직접 저장."""
    conn = sqlite3.connect(DB_PATH)
    now = now_str()
    created = 0

    for p in papers:
        cursor = conn.execute("""
            INSERT INTO papers (
                device_info_id, treatment_id, doi, title, title_ko,
                authors, journal, pub_year, pub_date,
                abstract_summary, key_findings, keywords,
                evidence_level, study_type, sample_size,
                source_url, source_file, status,
                one_line_summary, research_purpose, study_design_detail,
                key_results, conclusion, quotable_stats, cautions, follow_up_period,
                file_hash, created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            p.get("device_info_id"), p.get("treatment_id"), p.get("doi", ""),
            p["title"], p.get("title_ko", ""),
            p.get("authors", ""), p.get("journal", ""), p.get("pub_year"),
            p.get("pub_date", ""),
            p.get("abstract_summary", ""), p.get("key_findings", ""), p.get("keywords", "[]"),
            p.get("evidence_level", 0), p.get("study_type", ""), p.get("sample_size", ""),
            p.get("source_url", ""), p.get("source_file", ""), p.get("status", "draft"),
            p.get("one_line_summary", ""), p.get("research_purpose", ""),
            p.get("study_design_detail", ""), p.get("key_results", ""),
            p.get("conclusion", ""), p.get("quotable_stats", "[]"),
            p.get("cautions", ""), p.get("follow_up_period", ""),
            p.get("file_hash", ""), now, now,
        ))
        paper_id = cursor.lastrowid
        created += 1

        # paper_devices 다대다 연결 저장
        device_matches = p.get("_device_matches", [])
        for dm in device_matches:
            try:
                conn.execute("""
                    INSERT OR IGNORE INTO paper_devices
                        (paper_id, device_info_id, match_type, match_keyword, is_verified, created_at)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (
                    paper_id, dm["device_info_id"], dm["match_type"],
                    dm.get("match_keyword", ""), dm.get("is_verified", 1), now,
                ))
            except Exception as e:
                print(f"  [paper_devices 저장 오류] {e}")

        if device_matches:
            direct_count = sum(1 for m in device_matches if m["match_type"] == "direct")
            tech_count = sum(1 for m in device_matches if m["match_type"] == "technology")
            print(f"  → paper_devices: 직접 {direct_count}건 + 기술 {tech_count}건 연결")

    conn.commit()
    conn.close()
    print(f"\n[DB 저장 완료] {created}건 등록")
    return created


# ── Word 문서 출력 ──

def export_to_docx(papers: list[dict], output_path: str):
    """분석 결과를 Word 문서로 출력."""
    doc = Document()

    # 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.5

    # 제목
    title = doc.add_heading('논문 분석 결과 보고서', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"총 {len(papers)}건 분석")
    doc.add_page_break()

    # 시술별 그룹핑
    grouped = {}
    for p in papers:
        device = p.get("device_name") or "미분류"
        grouped.setdefault(device, []).append(p)

    for device_name, device_papers in grouped.items():
        doc.add_heading(f"관련 시술: {device_name}", level=1)

        for idx, p in enumerate(device_papers, 1):
            stars = EVIDENCE_STARS.get(p.get("evidence_level", 0), "☆☆☆☆☆")

            # 논문 헤더
            doc.add_heading(f"{idx}. {p.get('title_ko') or p.get('title', '제목 없음')}", level=2)

            # 기본 정보 테이블
            table = doc.add_table(rows=6, cols=2, style='Table Grid')
            table.columns[0].width = Cm(4)
            table.columns[1].width = Cm(12)

            info_rows = [
                ("원문 제목", p.get("title", "")),
                ("관련 시술", p.get("device_name") or "미분류"),
                ("근거 수준", f"{stars} ({p.get('evidence_level', 0)}/5) — {p.get('study_type', '')}"),
                ("대상자 / 추적기간", f"{p.get('sample_size', '미상')} / {p.get('follow_up_period', '미상')}"),
                ("저자 / 학술지", f"{p.get('authors', '')} / {p.get('journal', '')} ({p.get('pub_year', '')})"),
                ("DOI", p.get("doi", "")),
            ]
            for i, (label, value) in enumerate(info_rows):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
                # 레이블 셀 볼드
                for run in table.rows[i].cells[0].paragraphs[0].runs:
                    run.font.bold = True

            doc.add_paragraph()

            # 한 줄 요약
            doc.add_heading("한 줄 요약", level=3)
            doc.add_paragraph(p.get("one_line_summary", ""))

            # 핵심 결과
            doc.add_heading("핵심 결과", level=3)
            doc.add_paragraph(p.get("key_findings", ""))

            # 초록 기반 요약
            doc.add_heading("초록 기반 요약", level=2)

            doc.add_heading("1) 연구 목적", level=3)
            doc.add_paragraph(p.get("research_purpose", ""))

            doc.add_heading("2) 연구 설계 및 방법", level=3)
            doc.add_paragraph(p.get("study_design_detail", ""))

            doc.add_heading("3) 핵심 결과", level=3)
            doc.add_paragraph(p.get("key_results", ""))

            doc.add_heading("4) 연구 결론", level=3)
            doc.add_paragraph(p.get("conclusion", ""))

            # 해석 시 주의점
            doc.add_heading("해석 시 주의점", level=3)
            doc.add_paragraph(p.get("cautions", ""))

            doc.add_page_break()

    doc.save(output_path)
    print(f"[Word 출력] {output_path}")


# ── Excel 출력 ──

def export_to_xlsx(papers: list[dict], output_path: str):
    """분석 결과를 Excel로 출력."""
    wb = Workbook()
    # 기본 시트 삭제
    wb.remove(wb.active)

    # 시술별 그룹핑
    grouped = {}
    for p in papers:
        device = p.get("device_name") or "미분류"
        grouped.setdefault(device, []).append(p)

    header_font = Font(bold=True, color="FFFFFF", size=10)
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    wrap = Alignment(wrap_text=True, vertical="top")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin'),
    )

    headers = [
        "No", "제목 (한국어)", "원문 제목", "한 줄 요약", "핵심 결과",
        "근거 수준", "연구 유형", "대상자", "추적기간",
        "저자", "학술지", "출판연도", "DOI",
    ]

    for device_name, device_papers in grouped.items():
        # 시트명 (31자 제한, 특수문자 제거)
        sheet_name = re.sub(r'[\[\]:*?/\\]', '', device_name)[:31]
        ws = wb.create_sheet(title=sheet_name)

        # 헤더
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = wrap
            cell.border = thin_border

        # 데이터
        for row_idx, p in enumerate(device_papers, 2):
            stars = EVIDENCE_STARS.get(p.get("evidence_level", 0), "☆☆☆☆☆")
            values = [
                row_idx - 1,
                p.get("title_ko", ""),
                p.get("title", ""),
                p.get("one_line_summary", ""),
                p.get("key_findings", ""),
                f"{stars} ({p.get('evidence_level', 0)}/5)",
                p.get("study_type", ""),
                p.get("sample_size", ""),
                p.get("follow_up_period", ""),
                p.get("authors", ""),
                p.get("journal", ""),
                p.get("pub_year", ""),
                p.get("doi", ""),
            ]
            for col, val in enumerate(values, 1):
                cell = ws.cell(row=row_idx, column=col, value=val)
                cell.alignment = wrap
                cell.border = thin_border

        # 열 너비 설정
        widths = [5, 30, 30, 50, 40, 15, 15, 10, 10, 20, 20, 10, 25]
        for col, w in enumerate(widths, 1):
            ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = w

    wb.save(output_path)
    print(f"[Excel 출력] {output_path}")


# ── 콘솔 출력 ──

def _safe_print(text: str):
    """Windows 콘솔 인코딩 오류 방지 출력."""
    try:
        print(text)
    except UnicodeEncodeError:
        print(text.encode("utf-8", errors="replace").decode("utf-8"))


def print_result(p: dict):
    """분석 결과를 콘솔에 서술형으로 출력."""
    stars = EVIDENCE_STARS.get(p.get("evidence_level", 0), "☆☆☆☆☆")

    _safe_print(f"""
{'━'*60}
논문 분석 결과
{'━'*60}

관련 시술: {p.get('device_name') or '미분류'}
원문 제목: {p.get('title', '')}
근거 수준: {stars} ({p.get('evidence_level', 0)}/5) — {p.get('study_type', '')}

한 줄 요약:
{p.get('one_line_summary', '')}

핵심 결과:
{p.get('key_findings', '')}

{'─'*60}
초록 기반 요약
{'─'*60}

1) 연구 목적

{p.get('research_purpose', '')}

2) 연구 설계 및 방법

{p.get('study_design_detail', '')}

3) 핵심 결과

{p.get('key_results', '')}

4) 연구 결론

{p.get('conclusion', '')}

{'─'*60}
해석 시 주의점
{'─'*60}

{p.get('cautions', '')}

{'━'*60}
""")  # noqa: E501


# ── 메인 ──

def main():
    parser = argparse.ArgumentParser(description="미용의료 논문 PDF 분석 도구")
    parser.add_argument("files", nargs="*", help="분석할 PDF 파일(들)")
    parser.add_argument("--dir", help="PDF 파일이 들어있는 디렉토리")
    parser.add_argument("--dry-run", action="store_true", help="분석만 수행 (DB 저장 안 함)")
    parser.add_argument("--export-only", help="이미 분석된 JSON에서 문서만 출력")
    parser.add_argument("--api-url", default=API_BASE, help="uni-portal API 서버 URL")
    parser.add_argument("--export-docx", action="store_true", help="Word 문서도 출력")
    parser.add_argument("--export-xlsx", action="store_true", help="Excel 문서도 출력")
    parser.add_argument("--export-all", action="store_true", help="JSON + Word + Excel 모두 출력")
    args = parser.parse_args()

    os.makedirs(RESULTS_DIR, exist_ok=True)

    # export-all 플래그 처리
    want_docx = args.export_docx or args.export_all
    want_xlsx = args.export_xlsx or args.export_all
    want_json = args.export_all

    # export-only 모드
    if args.export_only:
        with open(args.export_only, "r", encoding="utf-8") as f:
            results = json.load(f)
        print(f"JSON에서 {len(results)}건 로드")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        export_to_docx(results, os.path.join(RESULTS_DIR, f"papers_{timestamp}.docx"))
        export_to_xlsx(results, os.path.join(RESULTS_DIR, f"papers_{timestamp}.xlsx"))
        return

    # PDF 파일 목록 수집
    pdf_files = []
    if args.dir:
        pdf_dir = Path(args.dir)
        pdf_files = sorted(pdf_dir.glob("*.pdf"))
    if args.files:
        pdf_files.extend([Path(f) for f in args.files])

    if not pdf_files:
        print("분석할 PDF 파일이 없습니다.")
        print("사용법: python paper_analyzer.py [--dir DIR] [files...]")
        sys.exit(1)

    # 존재하지 않는 파일 필터
    valid_files = []
    for f in pdf_files:
        if f.exists():
            valid_files.append(f)
        else:
            print(f"[경고] 파일 없음: {f}")
    pdf_files = valid_files

    if not pdf_files:
        print("유효한 PDF 파일이 없습니다.")
        sys.exit(1)

    print(f"\n총 {len(pdf_files)}건 PDF 분석 시작")
    print(f"{'='*60}")

    # API KEY 확인
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("[오류] ANTHROPIC_API_KEY 환경변수를 설정해주세요.")
        sys.exit(1)

    # DB 마이그레이션
    ensure_new_columns()

    # 장비/시술 목록 로드
    device_list = get_device_info_list()
    treatment_list = get_treatments_list()
    print(f"장비 사전: {len(device_list)}개 / 시술 사전: {len(treatment_list)}개")

    # 분석 실행
    results = []
    skipped = []
    failed = []
    for i, pdf_path in enumerate(pdf_files, 1):
        print(f"\n[{i}/{len(pdf_files)}]", end="")
        result = None
        try:
            result = process_pdf(str(pdf_path), device_list, treatment_list)
            if result and result.get("_skipped"):
                skipped.append(result)
            elif result:
                results.append(result)
                print_result(result)
            else:
                failed.append(str(pdf_path))
        except Exception as e:
            print(f"\n  [오류] {pdf_path}: {e}")
            failed.append(str(pdf_path))

        # API 속도 제한 대응 (1초 대기) — 중복 건은 API 미호출이므로 대기 불필요
        if i < len(pdf_files) and not (result and result.get("_skipped")):
            time.sleep(1)

    # 결과 저장
    if not results:
        if skipped:
            print(f"\n신규 분석 논문 없음 (중복 {len(skipped)}건 건너뜀)")
        else:
            print("\n분석 성공한 논문이 없습니다.")
        sys.exit(0 if skipped else 1)

    # 문서 출력 (요청 시에만)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    if want_json:
        json_path = os.path.join(RESULTS_DIR, f"papers_{timestamp}.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        print(f"\n[JSON 저장] {json_path}")
    if want_docx:
        export_to_docx(results, os.path.join(RESULTS_DIR, f"papers_{timestamp}.docx"))
    if want_xlsx:
        export_to_xlsx(results, os.path.join(RESULTS_DIR, f"papers_{timestamp}.xlsx"))

    # DB 저장
    if not args.dry_run:
        save_to_db(results)
    else:
        print("\n[dry-run 모드] DB 저장을 건너뛰었습니다.")

    # 비관련 논문 분류
    non_aesthetic = [r for r in results if r.get("_non_aesthetic")]
    unmatched = [r for r in results if not r.get("device_info_id") and not r.get("_non_aesthetic")]

    # 결과 요약
    print(f"\n{'='*60}")
    print(f"분석 완료")
    print(f"  성공: {len(results)}건")
    print(f"  중복 건너뜀: {len(skipped)}건")
    if skipped:
        for s in skipped:
            print(f"    - {Path(s['_file']).name}: {s['_reason']}")
    print(f"  실패: {len(failed)}건")
    if failed:
        for f_path in failed:
            print(f"    - {f_path}")

    # ⚠ 비관련 논문 알림
    if non_aesthetic:
        print(f"\n{'─'*60}")
        print(f"  ⚠ 미용의료 비관련 논문: {len(non_aesthetic)}건")
        print(f"  아래 논문은 미용의료와 관련 없는 것으로 판단되었습니다.")
        print(f"  DB에는 저장되었으나, 확인 후 삭제 또는 폴더에서 제거를 권장합니다.")
        print(f"{'─'*60}")
        for r in non_aesthetic:
            reason = r.get("_non_aesthetic_reason", "")
            print(f"    ❌ {r.get('title_ko', r.get('title', ''))}")
            print(f"       파일: {r.get('source_file', '')}")
            print(f"       사유: {reason}")

    # ⚠ 장비 미매칭 알림
    if unmatched:
        print(f"\n{'─'*60}")
        print(f"  ⚠ 장비 미매칭 논문: {len(unmatched)}건")
        print(f"  DB의 장비 목록과 매칭되지 않은 논문입니다. 장비 등록이 필요할 수 있습니다.")
        print(f"{'─'*60}")
        for r in unmatched:
            print(f"    ❓ {r.get('title_ko', r.get('title', ''))}")
            related = r.get("_related_devices_raw", [])
            if related:
                print(f"       논문에서 언급된 장비: {', '.join(related)}")

    print(f"{'='*60}")


if __name__ == "__main__":
    main()
